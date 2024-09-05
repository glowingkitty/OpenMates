import docx
import os

def print_document_structure(doc_path):
    doc = docx.Document(doc_path)
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: {paragraph.text}")
        for run in paragraph.runs:
            print(f"  Run: {run.text}")

def modify_document(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # Locate the "Professional Experience" section
    for i, paragraph in enumerate(doc.paragraphs):
        if "Professional Experience" in paragraph.text:
            # Add a new entry under "Professional Experience"
            new_entry = [
                "New Job Title at New Company",
                "New Company, Location",
                "Dates of Employment",
                "Description of responsibilities and achievements."
            ]
            for entry in new_entry:
                new_paragraph = doc.add_paragraph(entry)
                # Copy the style from the previous paragraph to maintain structure
                new_paragraph.style = paragraph.style
            break
    
    # Save the modified document
    doc.save(output_path)

# Example usage
script_dir = os.path.dirname(__file__)  # Get the directory of the script
doc_path = os.path.join(script_dir, 'Jane_Doe_CV.docx')  # Construct the full path to the document
output_path = os.path.join(script_dir, 'Jane_Doe_CV_Modified.docx')

print("Original Document Structure:")
print_document_structure(doc_path)

modify_document(doc_path, output_path)

print("\nModified Document Structure:")
print_document_structure(output_path)