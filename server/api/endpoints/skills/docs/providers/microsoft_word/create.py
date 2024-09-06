################
# Default Imports
################
import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO

# Fix import path
full_current_path = os.path.realpath(__file__)
main_directory = re.sub('server.*', '', full_current_path)
sys.path.append(main_directory)

from server import *
################

from typing import List
from server.api.models.skills.docs.skills_docs_create import DocsCreateInput, TextElement, HeadingElement, HyperlinkElement, ImageElement, TableElement, ListElement, CodeBlockElement, BlockQuoteElement, PageBreakElement

from server.api.models.skills.files.skills_files_upload import FilesUploadOutput
from server.api.endpoints.skills.files.upload import upload
from datetime import datetime, timedelta
import uuid


def add_text_element(doc, element: TextElement):
    p = doc.add_paragraph()
    run = p.add_run(element.text)
    if element.bold:
        run.bold = True
    if element.italic:
        run.italic = True
    if element.underline:
        run.underline = True
    if element.color:
        run.font.color.rgb = RGBColor.from_string(element.color[1:])
    if element.alignment:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, element.alignment.upper())
    if element.background_color:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), element.background_color[1:])
        p._element.get_or_add_pPr().append(shading_elm)

def add_heading_element(doc, element: HeadingElement):
    p = doc.add_heading(level=element.level)
    run = p.add_run(element.text)
    if element.color:
        run.font.color.rgb = RGBColor.from_string(element.color[1:])
    if element.alignment:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, element.alignment.upper())
    if element.background_color:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), element.background_color[1:])
        p._element.get_or_add_pPr().append(shading_elm)

def add_hyperlink_element(doc, element: HyperlinkElement):
    p = doc.add_paragraph()
    run = p.add_run(element.text)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True

    # Create the w:hyperlink tag and add needed values
    part = doc.part
    r_id = part.relate_to(element.url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:color element and set the value to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')

    # Create a w:u element and set the value to single (underline)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')

    # Add the w:color and w:u elements to the w:rPr element
    rPr.append(color)
    rPr.append(underline)

    # Add the w:rPr element to the w:r element
    new_run.append(rPr)

    # Create a w:t element and set the text value
    text_element = OxmlElement('w:t')
    text_element.text = element.text

    # Add the w:t element to the w:r element
    new_run.append(text_element)

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the w:hyperlink element to the paragraph
    p._element.append(hyperlink)

def add_image_element(doc, element: ImageElement):
    p = doc.add_paragraph()
    run = p.add_run()
    if element.url:
        # Download the image
        response = requests.get(element.url)
        if response.status_code == 200:
            image_stream = BytesIO(response.content)
            run.add_picture(image_stream, width=Pt(element.width.value) if element.width else None, height=Pt(element.height.value) if element.height else None)
        else:
            run.add_text(f"Image could not be downloaded: {element.url}")
    if element.alt_text:
        run.add_text(element.alt_text)
    if element.margin:
        p.paragraph_format.space_before = Pt(int(element.margin.split()[0].replace('px', '')))
        p.paragraph_format.space_after = Pt(int(element.margin.split()[1].replace('px', '')) if len(element.margin.split()) > 1 else 0)
    if element.border:
        p.paragraph_format.border = element.border

def add_table_element(doc, element: TableElement):
    table = doc.add_table(rows=len(element.rows), cols=len(element.rows[0].cells))
    for i, row in enumerate(element.rows):
        for j, cell in enumerate(row.cells):
            if hasattr(cell.content, 'text'):
                table.cell(i, j).text = cell.content.text
            else:
                table.cell(i, j).text = str(cell.content)  # Fallback to string representation
            if cell.background_color:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), cell.background_color[1:])
                table.cell(i, j)._element.get_or_add_tcPr().append(shading_elm)
            if cell.border:
                table.cell(i, j).paragraphs[0].paragraph_format.border = cell.border

def add_list_element(doc, element: ListElement):
    for item in element.items:
        p = doc.add_paragraph(style='ListNumber' if element.ordered else 'ListBullet')
        run = p.add_run(item.content.text)
        if item.level:
            p.paragraph_format.left_indent = Pt(item.level * 20)

def add_code_block_element(doc, element: CodeBlockElement):
    p = doc.add_paragraph()
    run = p.add_run(element.code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    if element.language:
        p.add_run(f' ({element.language})')

def add_block_quote_element(doc, element: BlockQuoteElement):
    p = doc.add_paragraph()
    run = p.add_run(f'"{element.text}"')
    run.italic = True
    if element.author:
        p.add_run(f' - {element.author}')

def add_page_break_element(doc, element: PageBreakElement):
    doc.add_page_break()

async def create(
    team_slug: str,
    api_token: str,
    title: str,
    elements: List[dict]
) -> FilesUploadOutput:
    """
    Create a new document
    """
    add_to_log(module_name="OpenMates | API | Microsoft Word | Create document", state="start", color="yellow", hide_variables=True)
    add_to_log("Creating a new document ...")

    doc = Document()
    for element in elements:
        element_type = element['type'] if type(element) == dict else element.type
        if element_type == 'text':
            add_text_element(doc, TextElement(**element)) if type(element) == dict else add_text_element(doc, element)
        elif element_type == 'heading':
            add_heading_element(doc, HeadingElement(**element)) if type(element) == dict else add_heading_element(doc, element)
        elif element_type == 'hyperlink':
            add_hyperlink_element(doc, HyperlinkElement(**element)) if type(element) == dict else add_hyperlink_element(doc, element)
        elif element_type == 'image':
            add_image_element(doc, ImageElement(**element)) if type(element) == dict else add_image_element(doc, element)
        elif element_type == 'table':
            add_table_element(doc, TableElement(**element)) if type(element) == dict else add_table_element(doc, element)
        elif element_type == 'list':
            add_list_element(doc, ListElement(**element)) if type(element) == dict else add_list_element(doc, element)
        elif element_type == 'code_block':
            add_code_block_element(doc, CodeBlockElement(**element)) if type(element) == dict else add_code_block_element(doc, element)
        elif element_type == 'block_quote':
            add_block_quote_element(doc, BlockQuoteElement(**element)) if type(element) == dict else add_block_quote_element(doc, element)
        elif element_type == 'page_break':
            add_page_break_element(doc, PageBreakElement(**element)) if type(element) == dict else add_page_break_element(doc, element)


    file_stream = BytesIO()
    doc.save(file_stream)
    file_data = file_stream.getvalue()

    file_name = f"{title.lower().replace(' ', '_')}.docx"
    file_id = uuid.uuid4().hex[:10]
    expiration_datetime = datetime.now() + timedelta(days=2)
    file_info = await upload(
        provider="docs",
        team_slug=team_slug,
        file_path=f"docs/{file_id}/{file_name}",
        name=file_name,
        file_data=file_data,
        file_id=file_id,
        encryption_key=api_token+file_id,
        expiration_datetime=expiration_datetime.isoformat(),
        access_public=False
    )
    # Clean up the file stream and data
    file_stream.close()
    del file_data

    return file_info