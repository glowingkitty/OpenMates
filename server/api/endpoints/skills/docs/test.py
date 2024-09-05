import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required URL
    :param text: The text displayed for the URL
    :return: The hyperlink object
    """
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:color element and set the value to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')

    # Create a w:u element and set the value to single (underline)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')

    # Add the w:color and w:u elements to the w:rPr element
    rPr.append(color)
    rPr.append(underline)

    # Add the w:rPr element to the w:r element
    new_run.append(rPr)

    # Create a w:t element and set the text value
    text_element = OxmlElement('w:t')
    text_element.text = text

    # Add the w:t element to the w:r element
    new_run.append(text_element)

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the w:hyperlink element to the paragraph
    paragraph._element.append(hyperlink)

    return hyperlink



def create_cv():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Add header
    header = doc.add_heading('Max Gulasch', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run('Berlin, Germany | Email: your.email@example.com | Phone: +49 123 456 7890')

    # Add professional summary
    doc.add_heading('Professional Summary', level=2)
    summary = doc.add_paragraph()
    summary.add_run('Skilled electronics enthusiast with experience in SMD soldering, PCB design, and prototyping. '
                    'Entrepreneurial background in developing RGB LED products. Seeking an SMD soldering position '
                    'to contribute expertise and further develop skills in specialized camera systems.')

    # Add skills section
    doc.add_heading('Skills', level=2)
    skills = [
        'SMD Soldering (hand soldering, microscope work, SMD oven, heat plate)',
        'PCB Design and Assembly',
        'Circuit Board Testing and Debugging',
        'Prototyping and Product Development',
        'Open Source Hardware Development',
        'Version Control (Git)',
        'CNC Machining, Laser Cutting, 3D Printing',
        'Software Development (Frontend, Backend, Embedded)'
    ]
    skill_para = doc.add_paragraph(style='List Bullet')
    for skill in skills:
        skill_para.add_run(skill + '\n')

    # Add experience section
    doc.add_heading('Relevant Experience', level=2)
    
    exp = doc.add_paragraph()
    exp.add_run('Founder and Electronics Developer').bold = True
    exp.add_run('\nOwn Business (Gewerbe), Berlin\n2021 - Present')
    doc.add_paragraph('• Design and develop various RGB LED products, including wands, wristbands, and lamps', style='List Bullet')
    doc.add_paragraph('• Create open-source hardware projects, sharing designs on GitHub', style='List Bullet')
    doc.add_paragraph('• Perform SMD soldering, PCB design, assembly, and debugging for prototypes', style='List Bullet')
    doc.add_paragraph('• Install custom RGB LED installations for clients in entertainment venues', style='List Bullet')

    # Add earlier work experience
    doc.add_heading('Additional Work Experience', level=2)
    
    experiences = [
        ('UX/UI Designer', 'Talentwunder GmbH, Berlin', '01.2016 - 04.2016'),
        ('UX Design Trainee', 'people interactive GmbH, Cologne', '07.2015 - 10.2015'),
        ('Junior Marketing Assistant', 'ASUS Computer GmbH, Ratingen', '10.2012 - 08.2013')
    ]
    
    for title, company, dates in experiences:
        exp = doc.add_paragraph()
        exp.add_run(f'{title} - {company}').bold = True
        exp.add_run(f'\n{dates}')

    # Add education section
    doc.add_heading('Education', level=2)
    
    edu = doc.add_paragraph()
    edu.add_run('FH Düsseldorf').bold = True
    edu.add_run('\n09.2013 - 08.2015')
    doc.add_paragraph('• Relevant coursework in electronics and design', style='List Bullet')

    edu = doc.add_paragraph()
    edu.add_run('Berufskolleg Opladen').bold = True
    edu.add_run('\nCompleted secondary education')

    # Add additional training
    doc.add_heading('Additional Training', level=2)
    training = doc.add_paragraph()
    training.add_run('Self-directed learning at various hackspaces and maker spaces worldwide')
    doc.add_paragraph('• Acquired skills in software development, electronics, PCB design, CNC machining, laser cutting, 3D printing, and more', style='List Bullet')

    # Save the document
    doc.save('Max_Gulasch_CV.docx')

create_cv()