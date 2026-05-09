# backend/apps/docs/tasks/generate_task.py
#
# Celery task for the Docs app artifact pipeline. It accepts the structured
# docx_model emitted by the LLM, creates a real DOCX file with python-docx,
# converts it to PDF via LibreOffice when available, renders page screenshots
# with PyMuPDF, then stores all artifacts encrypted in the private chatfiles
# bucket. The embed content keeps the structured model for future editing while
# browser preview/download use the generated artifacts.

import asyncio
import base64
import logging
import os
import shutil
import subprocess
import tempfile
import uuid
from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import fitz
import httpx
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from PIL import Image
from toon_format import decode as toon_decode
from toon_format import encode as toon_encode

from backend.core.api.app.services.embed_service import EmbedService
from backend.core.api.app.tasks.base_task import BaseServiceTask
from backend.core.api.app.tasks.celery_config import app

logger = logging.getLogger(__name__)

DOCS_S3_BUCKET = "chatfiles"
MAX_RETRIES = 2
DEFAULT_FILENAME = "OpenMates_Document.docx"


@app.task(
    bind=True,
    name="apps.docs.tasks.generate_docx",
    base=BaseServiceTask,
    queue="app_docs",
    soft_time_limit=300,
    time_limit=360,
    max_retries=MAX_RETRIES,
)
def generate_docx_task(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
    return asyncio.run(_async_generate_docx(self, arguments))


async def _async_generate_docx(task: BaseServiceTask, arguments: Dict[str, Any]) -> Dict[str, Any]:
    task_id = task.request.id
    embed_id: str = arguments["embed_id"]
    user_id: str = arguments["user_id"]
    user_id_hash: str = arguments["user_id_hash"]
    vault_key_id: str = arguments["vault_key_id"]
    chat_id: str = arguments["chat_id"]
    message_id: str = arguments["message_id"]
    docx_model: Dict[str, Any] = arguments["docx_model"]
    file_path_index: Dict[str, str] = arguments.get("file_path_index") or {}
    filename = _safe_docx_filename(arguments.get("filename") or docx_model.get("filename"))
    title = str(arguments.get("title") or docx_model.get("title") or "Document")
    log_prefix = f"[docs.generate] [task:{task_id[:8]}] [embed:{embed_id[:8]}]"
    created_s3_keys: List[str] = []

    try:
        await task.initialize_services()
        output_aes_key = os.urandom(32)
        output_aes_key_b64 = base64.b64encode(output_aes_key).decode("utf-8")
        output_aesgcm = AESGCM(output_aes_key)
        vault_wrapped_aes_key = await _wrap_key_via_vault(output_aes_key_b64, vault_key_id)

        with tempfile.TemporaryDirectory(prefix="openmates-docs-") as tmpdir:
            tmp_path = Path(tmpdir)
            docx_path = tmp_path / filename
            pdf_path = tmp_path / f"{docx_path.stem}.pdf"

            await _write_docx(task, docx_model, docx_path, title, file_path_index, vault_key_id, log_prefix)
            docx_bytes = docx_path.read_bytes()

            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            unique_id = uuid.uuid4().hex[:8]
            docx_s3_key = f"chatfiles/{user_id}/{timestamp}_{unique_id}_doc_{embed_id[:8]}.docx.bin"
            await _upload_encrypted(task, output_aesgcm, docx_s3_key, docx_bytes)
            created_s3_keys.append(docx_s3_key)

            screenshot_s3_keys: Dict[str, str] = {}
            page_count = 0
            if _convert_docx_to_pdf(docx_path, tmp_path, log_prefix) and pdf_path.exists():
                screenshots = _render_pdf_screenshots(pdf_path)
                page_count = len(screenshots)
                for page_num, png_bytes in sorted(screenshots.items()):
                    screenshot_key = (
                        f"chatfiles/{user_id}/{timestamp}_{unique_id}_doc_{embed_id[:8]}_p{page_num}.png.bin"
                    )
                    await _upload_encrypted(task, output_aesgcm, screenshot_key, png_bytes)
                    screenshot_s3_keys[str(page_num)] = screenshot_key
                    created_s3_keys.append(screenshot_key)
            else:
                logger.warning(f"{log_prefix} LibreOffice conversion unavailable; DOCX download will work without screenshots")

        word_count = _count_words(docx_model)
        embed_content = {
            "type": "document",
            "app_id": "docs",
            "skill_id": "document",
            "status": "finished",
            "title": title,
            "filename": filename,
            "word_count": word_count,
            "page_count": page_count,
            "docx_model": docx_model,
            "docx_s3_key": docx_s3_key,
            "screenshot_s3_keys": screenshot_s3_keys,
            "aes_key": output_aes_key_b64,
            "aes_nonce": "",
            "vault_wrapped_aes_key": vault_wrapped_aes_key,
        }
        content_toon = toon_encode(embed_content)
        encrypted_toon, _ = await task._encryption_service.encrypt_with_user_key(content_toon, vault_key_id)
        now_ts = int(datetime.now(timezone.utc).timestamp())
        embed_service = EmbedService(task._cache_service, task._directus_service, task._encryption_service)
        cached_embed = await embed_service._get_cached_embed(embed_id, vault_key_id, log_prefix) or {}
        updated_embed = {
            **cached_embed,
            "embed_id": embed_id,
            "type": "document",
            "chat_id": chat_id,
            "message_id": message_id,
            "status": "finished",
            "hashed_user_id": user_id_hash,
            "encryption_mode": "client",
            "encrypted_content": encrypted_toon,
            "updated_at": now_ts,
            "created_at": cached_embed.get("created_at") or now_ts,
        }
        await embed_service._cache_embed(embed_id, updated_embed, chat_id, user_id_hash, vault_key_id, user_id)
        await embed_service.send_embed_data_to_client(
            embed_id=embed_id,
            embed_type="document",
            content_toon=content_toon,
            chat_id=chat_id,
            message_id=message_id,
            user_id=user_id,
            user_id_hash=user_id_hash,
            status="finished",
            encryption_mode="client",
            created_at=updated_embed["created_at"],
            updated_at=now_ts,
            log_prefix=log_prefix,
            check_cache_status=False,
        )
        embed_service._schedule_embed_persistence_fallback(embed_id)
        logger.info(f"{log_prefix} DOCX generation complete: pages={page_count}, words={word_count}")
        return {"embed_id": embed_id, "status": "finished", "page_count": page_count}

    except Exception as exc:
        logger.error(f"{log_prefix} DOCX generation failed: {exc}", exc_info=True)
        await _mark_embed_error(task, arguments, str(exc), log_prefix)
        raise


async def _write_docx(
    task: BaseServiceTask,
    model: Dict[str, Any],
    path: Path,
    fallback_title: str,
    file_path_index: Dict[str, str],
    vault_key_id: str,
    log_prefix: str,
) -> None:
    document = Document()
    core = document.core_properties
    core.title = fallback_title
    for block in model.get("blocks", []):
        block_type = block.get("type")
        if block_type == "heading":
            level = int(block.get("level") or 1)
            paragraph = document.add_heading(str(block.get("text") or ""), level=max(1, min(level, 4)))
            _apply_alignment(paragraph, block.get("align"))
        elif block_type == "paragraph":
            paragraph = document.add_paragraph()
            _add_runs(paragraph, block)
            _apply_alignment(paragraph, block.get("align"))
        elif block_type == "list":
            style = "List Number" if block.get("ordered") else "List Bullet"
            for item in block.get("items", []):
                document.add_paragraph(str(item), style=style)
        elif block_type == "table":
            rows = block.get("rows") or []
            headers = block.get("headers") or []
            if headers or rows:
                table = document.add_table(rows=1 if headers else 0, cols=max(len(headers), len(rows[0]) if rows else 1))
                table.style = "Table Grid"
                if headers:
                    for idx, header in enumerate(headers):
                        table.rows[0].cells[idx].text = str(header)
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row[: len(cells)]):
                        cells[idx].text = str(value)
        elif block_type == "blockquote":
            paragraph = document.add_paragraph(str(block.get("text") or ""), style="Intense Quote")
        elif block_type == "image":
            embed_ref = block.get("embed_ref")
            if not isinstance(embed_ref, str) or not embed_ref:
                continue
            image_bytes = await _resolve_image_bytes(task, embed_ref, file_path_index, vault_key_id, log_prefix)
            if image_bytes:
                image_stream = BytesIO(_normalize_image_for_docx(image_bytes))
                width_inches = float(block.get("width_inches") or 5.5)
                document.add_picture(image_stream, width=Inches(max(1.0, min(width_inches, 6.5))))
        elif block_type == "page_break":
            document.add_page_break()
    document.save(path)


def _add_runs(paragraph: Any, block: Dict[str, Any]) -> None:
    runs = block.get("runs")
    if not isinstance(runs, list):
        runs = [{"text": block.get("text") or ""}]
    for run_data in runs:
        run = paragraph.add_run(str(run_data.get("text") or ""))
        run.bold = bool(run_data.get("bold"))
        run.italic = bool(run_data.get("italic"))
        run.underline = bool(run_data.get("underline"))
        color = run_data.get("color")
        if isinstance(color, str) and color.startswith("#") and len(color) == 7:
            run.font.color.rgb = RGBColor.from_string(color[1:])


def _apply_alignment(paragraph: Any, align: Optional[str]) -> None:
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _convert_docx_to_pdf(docx_path: Path, output_dir: Path, log_prefix: str) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=120,
    )
    if result.returncode != 0:
        logger.warning(f"{log_prefix} LibreOffice conversion failed: {result.stderr.decode('utf-8', 'ignore')[:400]}")
        return False
    return True


def _render_pdf_screenshots(pdf_path: Path) -> Dict[int, bytes]:
    screenshots: Dict[int, bytes] = {}
    pdf = fitz.open(pdf_path)
    try:
        matrix = fitz.Matrix(150 / 72, 150 / 72)
        for index, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            screenshots[index] = pix.tobytes("png")
    finally:
        pdf.close()
    return screenshots


async def _upload_encrypted(task: BaseServiceTask, aesgcm: AESGCM, s3_key: str, content: bytes) -> None:
    nonce = os.urandom(12)
    encrypted = nonce + aesgcm.encrypt(nonce, content, None)
    await task._s3_service.upload_file(
        bucket_key=DOCS_S3_BUCKET,
        file_key=s3_key,
        content=encrypted,
        content_type="application/octet-stream",
    )


async def _wrap_key_via_vault(aes_key_b64: str, vault_key_id: str) -> str:
    vault_url = os.environ.get("VAULT_URL", "http://vault:8200")
    with open("/vault-data/api.token") as token_file:
        vault_token = token_file.read().strip()
    context = base64.b64encode(vault_key_id.encode()).decode("utf-8")
    encoded_plaintext = base64.b64encode(aes_key_b64.encode()).decode("utf-8")
    async with httpx.AsyncClient(timeout=15) as client:
        response = await client.post(
            f"{vault_url}/v1/transit/encrypt/{vault_key_id}",
            json={"plaintext": encoded_plaintext, "context": context},
            headers={"X-Vault-Token": vault_token},
        )
    if response.status_code != 200:
        raise RuntimeError(f"Vault transit encrypt failed: HTTP {response.status_code} - {response.text[:200]}")
    return response.json()["data"]["ciphertext"]


async def _unwrap_key_via_vault(vault_wrapped_aes_key: str, vault_key_id: str) -> bytes:
    vault_url = os.environ.get("VAULT_URL", "http://vault:8200")
    with open("/vault-data/api.token") as token_file:
        vault_token = token_file.read().strip()
    context = base64.b64encode(vault_key_id.encode()).decode("utf-8")
    async with httpx.AsyncClient(timeout=15) as client:
        response = await client.post(
            f"{vault_url}/v1/transit/decrypt/{vault_key_id}",
            json={"ciphertext": vault_wrapped_aes_key, "context": context},
            headers={"X-Vault-Token": vault_token},
        )
    if response.status_code != 200:
        raise RuntimeError(f"Vault transit decrypt failed: HTTP {response.status_code} - {response.text[:200]}")
    aes_key_b64 = base64.b64decode(response.json()["data"]["plaintext"]).decode("utf-8")
    return base64.b64decode(aes_key_b64)


async def _resolve_image_bytes(
    task: BaseServiceTask,
    embed_ref: str,
    file_path_index: Dict[str, str],
    vault_key_id: str,
    log_prefix: str,
) -> Optional[bytes]:
    embed_id = file_path_index.get(embed_ref)
    if not embed_id:
        logger.warning(f"{log_prefix} Image embed_ref {embed_ref!r} was not found in file_path_index")
        return None
    cached_embed = await task._cache_service.get_embed_from_cache(embed_id)
    if not cached_embed or not cached_embed.get("encrypted_content"):
        logger.warning(f"{log_prefix} Image embed {embed_id[:8]} missing encrypted cache content")
        return None
    decrypted_toon = await task._encryption_service.decrypt_with_user_key(
        cached_embed["encrypted_content"], vault_key_id
    )
    embed_content = toon_decode(decrypted_toon)
    if not isinstance(embed_content, dict):
        return None

    image_url = embed_content.get("image_url")
    if isinstance(image_url, str) and image_url.startswith(("http://", "https://")):
        async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
            response = await client.get(image_url)
        if response.status_code == 200:
            return response.content
        logger.warning(f"{log_prefix} Failed to fetch image_url for {embed_ref}: HTTP {response.status_code}")
        return None

    vault_wrapped_aes_key = embed_content.get("vault_wrapped_aes_key")
    aes_nonce = embed_content.get("aes_nonce")
    files = embed_content.get("files") or {}
    if not vault_wrapped_aes_key or not aes_nonce or not isinstance(files, dict):
        logger.warning(f"{log_prefix} Image embed {embed_ref} missing encrypted file metadata")
        return None
    s3_key = None
    for variant_name in ("full", "original", "preview"):
        variant = files.get(variant_name)
        if isinstance(variant, dict) and variant.get("s3_key"):
            s3_key = variant["s3_key"]
            break
    if not s3_key:
        return None
    download_url = f"{os.environ.get('INTERNAL_API_BASE_URL', 'http://api:8000')}/internal/s3/download"
    shared_token = os.environ.get("INTERNAL_API_SHARED_TOKEN", "")
    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.get(
            download_url,
            params={"bucket_key": "chatfiles", "s3_key": s3_key},
            headers={"X-Internal-Service-Token": shared_token},
        )
    if response.status_code != 200:
        logger.warning(f"{log_prefix} Failed to download image {s3_key}: HTTP {response.status_code}")
        return None
    image_aes_key = await _unwrap_key_via_vault(vault_wrapped_aes_key, vault_key_id)
    return AESGCM(image_aes_key).decrypt(base64.b64decode(aes_nonce), response.content, None)


def _normalize_image_for_docx(image_bytes: bytes) -> bytes:
    with Image.open(BytesIO(image_bytes)) as image:
        if image.format in {"PNG", "JPEG"}:
            return image_bytes
        output = BytesIO()
        image.convert("RGBA").save(output, format="PNG")
        return output.getvalue()


async def _mark_embed_error(task: BaseServiceTask, arguments: Dict[str, Any], error: str, log_prefix: str) -> None:
    try:
        embed_service = EmbedService(task._cache_service, task._directus_service, task._encryption_service)
        content_toon = toon_encode({
            "type": "document",
            "app_id": "docs",
            "skill_id": "document",
            "status": "error",
            "title": arguments.get("title") or "Document",
            "filename": _safe_docx_filename(arguments.get("filename")),
            "error": error,
            "docx_model": arguments.get("docx_model") or {},
        })
        await embed_service.send_embed_data_to_client(
            embed_id=arguments["embed_id"],
            embed_type="document",
            content_toon=content_toon,
            chat_id=arguments["chat_id"],
            message_id=arguments["message_id"],
            user_id=arguments["user_id"],
            user_id_hash=arguments["user_id_hash"],
            status="error",
            encryption_mode="client",
            log_prefix=log_prefix,
            check_cache_status=False,
        )
    except Exception:
        logger.error(f"{log_prefix} Failed to mark document embed as error", exc_info=True)


def _safe_docx_filename(value: Any) -> str:
    raw = str(value or DEFAULT_FILENAME).strip() or DEFAULT_FILENAME
    name = Path(raw).name.replace(" ", "_")
    safe = "".join(ch for ch in name if ch.isalnum() or ch in ("-", "_", "."))
    if not safe.lower().endswith(".docx"):
        safe = f"{safe or 'OpenMates_Document'}.docx"
    stem = safe[:-5] or "OpenMates_Document"
    return f"{stem[:115]}.docx"


def _count_words(model: Dict[str, Any]) -> int:
    parts: List[str] = []
    for block in model.get("blocks", []):
        for key in ("text", "items", "headers", "rows"):
            value = block.get(key)
            if isinstance(value, str):
                parts.append(value)
            elif isinstance(value, list):
                parts.append(" ".join(str(item) for item in value))
    return len(" ".join(parts).split())
